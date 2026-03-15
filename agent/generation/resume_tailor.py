"""Resume tailoring: clones master_resume.docx, injects keywords, optionally reorders sections.

Rules (from SPEC.md):
- Keywords are injected into existing content — never fabricated
- Only the Summary paragraph and Skills section are modified
- Section reorder: changes top-level section order only; content within sections is untouched
- The master resume is never written to — always shutil.copy'd first
"""
import json
import logging
import re
import shutil
from pathlib import Path
from typing import Any

from docx import Document  # type: ignore
from docx.oxml.ns import qn  # type: ignore

from agent.db.models import Job
from agent.scoring.llm_scorer import LLMBackend
from agent.scoring.prompts import RESUME_TAILOR_SYSTEM_PROMPT, RESUME_TAILOR_USER_PROMPT

logger = logging.getLogger(__name__)


def _docx_full_text(doc: Document) -> str:
    """Extract all paragraph text from a docx Document."""
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def _find_section_paragraph(doc: Document, heading_pattern: str) -> int | None:
    """Return paragraph index of the first heading matching pattern (case-insensitive)."""
    pattern = re.compile(heading_pattern, re.IGNORECASE)
    for i, para in enumerate(doc.paragraphs):
        if para.style.name.startswith("Heading") and pattern.search(para.text):
            return i
    return None


def _inject_into_paragraph(para, keywords: list[str]) -> None:
    """Append comma-separated keywords to the first run of a paragraph."""
    if not keywords:
        return
    addition = ", ".join(keywords)
    if para.runs:
        para.runs[-1].text = para.runs[-1].text.rstrip() + f", {addition}"
    else:
        para.add_run(f" {addition}")


def tailor_resume(
    master_resume_path: str,
    output_path: str,
    job: Job,
    profile_text: str,
    llm: LLMBackend,
) -> dict[str, Any]:
    """
    Clone master resume → output_path, then apply LLM-driven tailoring.
    Returns a dict of changes made (for notes.md).
    Never modifies master_resume_path.
    """
    # Safety guard — should never happen given orchestrator logic, but belt+suspenders
    master = Path(master_resume_path).resolve()
    output = Path(output_path).resolve()
    if master == output:
        raise ValueError("Output path must differ from master resume path")

    shutil.copy(str(master), str(output))
    doc = Document(str(output))
    resume_text = _docx_full_text(doc)

    # Ask LLM for tailoring instructions
    system_prompt = RESUME_TAILOR_SYSTEM_PROMPT
    user_prompt = RESUME_TAILOR_USER_PROMPT.format(
        profile_text=profile_text,
        job_description=f"{job.title}\n{job.company}\n\n{job.description or ''}",
        resume_text=resume_text,
    )

    changes: dict[str, Any] = {
        "summary_additions": [],
        "skills_keywords": [],
        "section_reorder": [],
        "mapping_notes": [],
    }

    try:
        raw = llm.complete(system_prompt, user_prompt)
        data = json.loads(raw)
        changes.update(data)
    except Exception as exc:
        logger.warning("Resume tailor LLM call failed: %s", exc)
        doc.save(str(output))
        return changes

    # Apply summary additions
    summary_additions: list[str] = changes.get("summary_additions", [])
    if summary_additions:
        summary_idx = _find_section_paragraph(doc, r"summary|objective|profile")
        if summary_idx is not None and summary_idx + 1 < len(doc.paragraphs):
            para = doc.paragraphs[summary_idx + 1]
            if para.runs:
                existing = para.runs[-1].text.rstrip()
                addition = " " + "; ".join(summary_additions)
                if addition.rstrip(". ;") not in existing:
                    para.runs[-1].text = existing + addition

    # Apply skills keywords
    skills_keywords: list[str] = changes.get("skills_keywords", [])
    if skills_keywords:
        skills_idx = _find_section_paragraph(doc, r"skills|technologies|technical")
        if skills_idx is not None and skills_idx + 1 < len(doc.paragraphs):
            para = doc.paragraphs[skills_idx + 1]
            _inject_into_paragraph(para, skills_keywords)

    # Apply section reorder (move XML elements)
    section_order: list[str] = changes.get("section_reorder", [])
    if section_order:
        _reorder_sections(doc, section_order)

    doc.save(str(output))
    logger.info("Resume tailored → %s", output)
    return changes


def _reorder_sections(doc: Document, desired_order: list[str]) -> None:
    """Reorder top-level heading sections in the docx body XML."""
    body = doc.element.body
    children = list(body)

    # Map heading text → list of element indices (heading + all following paragraphs until next heading)
    sections: list[tuple[str, list]] = []
    current_heading: str | None = None
    current_elements: list = []

    for child in children:
        # Check if this is a heading paragraph
        pPr = child.find(qn("w:pPr"))
        pStyle = pPr.find(qn("w:pStyle")) if pPr is not None else None
        style_val = pStyle.get(qn("w:val"), "") if pStyle is not None else ""

        if style_val.lower().startswith("heading"):
            if current_heading is not None:
                sections.append((current_heading, current_elements))
            # Get text of this heading
            texts = [t.text for t in child.iter(qn("w:t")) if t.text]
            current_heading = "".join(texts)
            current_elements = [child]
        else:
            current_elements.append(child)

    if current_heading is not None:
        sections.append((current_heading, current_elements))

    # Build a mapping from normalized section name to section data
    section_map = {name.lower(): (name, elements) for name, elements in sections}

    # Determine reorder: only move sections that appear in desired_order
    reordered_names = [n.lower() for n in desired_order]
    all_names = [name.lower() for name, _ in sections]
    remaining = [n for n in all_names if n not in reordered_names]
    final_order = reordered_names + remaining

    # Rebuild body
    # Remove all section children first
    for _, elems in sections:
        for el in elems:
            body.remove(el)

    for name in final_order:
        if name in section_map:
            _, elems = section_map[name]
            for el in elems:
                body.append(el)
