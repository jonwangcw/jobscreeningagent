"""Cover letter generation: LLM call → writes cover_letter.docx."""
import logging
from pathlib import Path

from docx import Document  # type: ignore
from docx.shared import Pt  # type: ignore

from agent.db.models import Job
from agent.scoring.llm_scorer import LLMBackend
from agent.scoring.prompts import COVER_LETTER_SYSTEM_PROMPT, COVER_LETTER_USER_PROMPT

logger = logging.getLogger(__name__)


def generate_cover_letter(
    output_path: str,
    job: Job,
    profile_text: str,
    llm: LLMBackend,
) -> None:
    """Generate a cover letter docx at output_path."""
    user_prompt = COVER_LETTER_USER_PROMPT.format(
        profile_text=profile_text,
        title=job.title,
        company=job.company,
        description=(job.description or "")[:4000],
    )

    try:
        letter_text = llm.complete(COVER_LETTER_SYSTEM_PROMPT, user_prompt)
    except Exception as exc:
        logger.error("Cover letter LLM call failed: %s", exc)
        letter_text = f"[Cover letter generation failed: {exc}]"

    doc = Document()
    doc.add_paragraph()  # spacer

    for paragraph_text in letter_text.strip().split("\n\n"):
        para = doc.add_paragraph(paragraph_text.strip())
        para.style.font.size = Pt(11)

    doc.save(output_path)
    logger.info("Cover letter written → %s", output_path)
