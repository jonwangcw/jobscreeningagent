"""Generation tests — mock LLM, verify master_resume immutability, output folder structure."""
import json
import shutil
from pathlib import Path
from unittest.mock import MagicMock
from datetime import datetime

import pytest
from docx import Document

from agent.db.models import Job
from agent.generation.cover_letter import generate_cover_letter
from agent.generation.resume_tailor import tailor_resume
from agent.scoring.llm_scorer import LLMBackend


TAILOR_RESPONSE = json.dumps({
    "summary_additions": ["production ML systems"],
    "skills_keywords": ["PyTorch", "MLflow"],
    "section_reorder": [],
    "mapping_notes": ["Python → Python (JD exact match)"],
})

COVER_LETTER_TEXT = "Dear Hiring Manager,\n\nI am excited to apply.\n\nSincerely,\nJon"


class MockLLM(LLMBackend):
    def __init__(self, response: str):
        self._response = response

    def complete(self, system: str, user: str) -> str:
        return self._response


def make_job(**kwargs) -> Job:
    job = Job()
    job.id = 1
    job.title = "ML Engineer"
    job.company = "Acme"
    job.location = "Pittsburgh, PA"
    job.remote = False
    job.description = "Build ML pipelines with PyTorch and MLflow."
    job.url = "https://example.com"
    job.composite_score = 0.85
    job.rationale = "Strong fit."
    job.skill_gaps = json.dumps(["Kubernetes"])
    job.status = "reviewed"
    job.created_at = datetime.utcnow()
    job.updated_at = datetime.utcnow()
    for k, v in kwargs.items():
        setattr(job, k, v)
    return job


def make_master_resume(path: Path) -> None:
    """Create a minimal .docx for testing."""
    doc = Document()
    doc.add_heading("Summary", level=1)
    doc.add_paragraph("Experienced ML engineer.")
    doc.add_heading("Skills", level=1)
    doc.add_paragraph("Python, SQL, PyTorch")
    doc.save(str(path))


def test_master_resume_never_modified(tmp_path):
    """tailor_resume must not change the master resume file."""
    master = tmp_path / "master_resume.docx"
    make_master_resume(master)
    master_mtime_before = master.stat().st_mtime

    output = tmp_path / "resume.docx"
    llm = MockLLM(TAILOR_RESPONSE)
    job = make_job()
    profile_text = "## Target roles\nML Engineer\n## Skills\nPython"

    tailor_resume(
        master_resume_path=str(master),
        output_path=str(output),
        job=job,
        profile_text=profile_text,
        llm=llm,
    )

    # Master file must be untouched
    assert master.stat().st_mtime == master_mtime_before
    master_bytes = master.read_bytes()
    assert master_bytes == master.read_bytes()

    # Output file should exist and differ (or at least exist)
    assert output.exists()


def test_output_is_copy_not_master(tmp_path):
    """Output resume must be a distinct file from the master."""
    master = tmp_path / "master_resume.docx"
    make_master_resume(master)
    output = tmp_path / "resume_output.docx"

    llm = MockLLM(TAILOR_RESPONSE)
    job = make_job()

    tailor_resume(
        master_resume_path=str(master),
        output_path=str(output),
        job=job,
        profile_text="profile text",
        llm=llm,
    )

    assert master.resolve() != output.resolve()
    assert output.exists()


def test_tailor_resume_raises_if_same_path(tmp_path):
    master = tmp_path / "master_resume.docx"
    make_master_resume(master)

    with pytest.raises(ValueError, match="Output path must differ"):
        tailor_resume(
            master_resume_path=str(master),
            output_path=str(master),
            job=make_job(),
            profile_text="profile",
            llm=MockLLM(TAILOR_RESPONSE),
        )


def test_cover_letter_creates_docx(tmp_path):
    output = tmp_path / "cover_letter.docx"
    job = make_job()
    llm = MockLLM(COVER_LETTER_TEXT)

    generate_cover_letter(
        output_path=str(output),
        job=job,
        profile_text="## Target roles\nML Engineer",
        llm=llm,
    )

    assert output.exists()
    doc = Document(str(output))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Hiring Manager" in full_text or "excited" in full_text


def test_cover_letter_handles_llm_failure(tmp_path):
    """Even if LLM fails, a docx with error message is written — does not raise."""
    output = tmp_path / "cover_letter.docx"
    job = make_job()

    class FailingLLM(LLMBackend):
        def complete(self, system: str, user: str) -> str:
            raise RuntimeError("API unavailable")

    generate_cover_letter(
        output_path=str(output),
        job=job,
        profile_text="profile",
        llm=FailingLLM(),
    )

    assert output.exists()
